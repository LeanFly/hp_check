from docx import Document

def generate_report(name, gender, age, department, patient_id, clinical_diagnosis, referring_doctor):
    doc = Document()

    doc.add_heading('复旦大学附属闵行医院/复旦大学附属中山医院闵行分院', level=1)
    doc.add_paragraph('上海市闵行区中心医院')
    doc.add_paragraph('幽门螺杆菌快速检测报告单')

    info = f"""
姓名: {name}       性别：{gender}      年龄：{age}      标本种类：末梢血    送检日期: 

科室：{department}       门诊号/住院号：{patient_id}          临床诊断：{clinical_diagnosis}          送检医生：{referring_doctor}
    """
    doc.add_paragraph(info)

    doc.add_paragraph('序号', style='List Number')
    doc.add_paragraph('检查项目', style='List Number')
    doc.add_paragraph('结果', style='List Number')
    doc.add_paragraph('参考值', style='List Number')

    doc.add_paragraph('1', style='List Number')
    doc.add_paragraph('HP-CIMHP-现感染', style='List Number')
    doc.add_paragraph('阳性（+）', style='List Number')
    doc.add_paragraph('阴性（—）', style='List Number')

    doc.add_paragraph('2', style='List Number')
    doc.add_paragraph('HP-IgGHP-IgG', style='List Number')
    doc.add_paragraph('阳性（+）', style='List Number')
    doc.add_paragraph('阴性（—）', style='List Number')

    doc.add_paragraph('检验时间：', style='List Number')
    doc.add_paragraph('报告时间：', style='List Number')
    doc.add_paragraph('检验者：', style='List Number')
    doc.add_paragraph('审核者：', style='List Number')

    doc.add_paragraph('项目前带*表示已复做，此结果仅对该标本负责，且仅供本院医师参考。')

    doc.save(f'{name}_HP_Report.docx')

if __name__ == "__main__":
    name = input("请输入姓名：")
    gender = input("请输入性别：")
    age = input("请输入年龄：")
    department = input("请输入科室：")
    patient_id = input("请输入门诊号/住院号：")
    clinical_diagnosis = input("请输入临床诊断：")
    referring_doctor = input("请输入送检医生：")

    generate_report(name, gender, age, department, patient_id, clinical_diagnosis, referring_doctor)