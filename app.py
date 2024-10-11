# 构建 doc 文档的包
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_UNDERLINE
import re
from datetime import datetime


# 基准文档
doc = Document("HP1.docx")

def build_doc():
    # 修改报告里的表格数据。
    today = datetime.today()
    date_now = today.strftime('%Y-%m-%d')
    
    tables = doc.tables
    
    # 处理第一行的姓名、性别、年龄、标本种类、送检日期
    table1 = tables[0]
    run = table1.cell(0, 0).paragraphs[0].add_run() # 姓名
    run.text = "赵某某"
    run = table1.cell(0, 1).paragraphs[0].add_run() # 行别
    run.text = "男"
    run = table1.cell(0, 2).paragraphs[0].add_run() # 年龄
    run.text = "33"
    run = table1.cell(0, 3).paragraphs[0].add_run() # 类型
    run.text = "末梢血"
    run = table1.cell(0, 4).paragraphs[0].add_run() # 日期
    run.text = f"{date_now}"
    
    # 处理第二行的科室、门诊号/住院号、临床诊断、送检医生
    table2 = tables[1]
    run = table2.cell(0, 0).paragraphs[0].add_run() # 科室
    run.text = "检验科"
    run = table2.cell(0, 1).paragraphs[0].add_run() # 门诊号
    run.text = "99"
    run = table2.cell(0, 2).paragraphs[0].add_run() # 诊断结果
    run.text = "过敏性皮炎"
    run = table2.cell(0, 3).paragraphs[0].add_run() # 医生
    run.text = "值班医生"
    
    
    # 处理检测结果序号、检查项目、结果、参考值
    table4 = tables[2]
    run = table4.cell(1, 0).paragraphs[0].add_run() # 序号
    run.text = "1"
    run = table4.cell(1, 1).paragraphs[0].add_run() # 项目
    run.text = "HP-CIMHP-现感染"
    run = table4.cell(1, 2).paragraphs[0].add_run() # 结果
    run.text = "阴性（—）"
    run = table4.cell(1, 3).paragraphs[0].add_run() # 参考值
    run.text = "阴性（—）"
    
    run = table4.cell(2, 0).paragraphs[0].add_run() # 序号
    run.text = "2"
    run = table4.cell(2, 1).paragraphs[0].add_run() # 项目
    run.text = "HP-CIMHP-现感染"
    run = table4.cell(2, 2).paragraphs[0].add_run() # 结果
    run.text = "阴性（—）"
    run = table4.cell(2, 3).paragraphs[0].add_run() # 参考值
    run.text = "阴性（—）"
    
    
    
    
    # 处理底部的检验时间、报告时间、检验者、审核者
    # table_list[1].rows[0].cells[0].text   '检验时间：                     报告时间：            检验者：       审核者：                 '
    table_l = tables[-1]
    
    run = table_l.cell(0, 0).paragraphs[0].add_run()
    run.text = f"{date_now}"
    run = table_l.cell(0, 1).paragraphs[0].add_run()
    run.text = f"{date_now}"
    
    run = table_l.cell(0, 2).paragraphs[0].add_run()
    run.text = f"理理理"
    run = table_l.cell(0, 3).paragraphs[0].add_run()
    run.text = f"理理理"
    
    
    doc.save(f"test_{int(today.timestamp())}.docx")
    
    
    
if __name__ == "__main__":
    
    build_doc()