from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def generate_report(name, gender, age, department, patient_id, clinical_diagnosis, referring_doctor, test_date, report_date):
    doc = Document()

    # 添加医院信息
    header = doc.add_heading(level=1)
    header.text = '复旦大学附属闵行医院/复旦大学附属中山医院闵行分院'
    header.alignment = 1  # 居中对齐

    doc.add_paragraph('上海市闵行区中心医院')
    doc.add_paragraph('幽门螺杆菌快速检测报告单')
    doc.add_paragraph()

    # 添加患者信息
    patient_info = f"姓名: {name}\n性别：{gender}\n年龄：{age}\n标本种类：末梢血\n送检日期: \n科室：{department}\n门诊号/住院号：{patient_id}\n临床诊断：{clinical_diagnosis}\n送检医生：{referring_doctor}"
    doc.add_paragraph(patient_info, style='List Bullet')

    # 添加表格
    table = doc.add_table(rows=3, cols=4)
    table.style = 'Table Grid'

    # 设置表头
    cells = table.rows[0].cells
    cells[0].text = '序号'
    cells[1].text = '检查项目'
    cells[2].text = '结果'
    cells[3].text = '参考值'

    # 设置第一行数据
    cells = table.rows[1].cells
    cells[0].text = '1'
    cells[1].text = 'HP-CIMHP-现感染'
    cells[2].text = '阳性（+）'
    cells[3].text = '阴性（—）'

    # 设置第二行数据
    cells = table.rows[2].cells
    cells[0].text = '2'
    cells[1].text = 'HP-IgGHP-IgG'
    cells[2].text = '阳性（+）'
    cells[3].text = '阴性（—）'

    # 添加检验时间和报告时间
    doc.add_paragraph()
    doc.add_paragraph(f"检验时间：{test_date}")
    doc.add_paragraph(f"报告时间：{report_date}")
    doc.add_paragraph("检验者：")
    doc.add_paragraph("审核者：")

    # 添加备注
    doc.add_paragraph("项目前带*表示已复做，此结果仅对该标本负责，且仅供本院医师参考。")

    # 保存文档
    doc.save(f'{name}_HP_Report.docx')

if __name__ == "__main__":
    name = input("请输入姓名：")
    gender = input("请输入性别：")
    age = input("请输入年龄：")
    department = input("请输入科室：")
    patient_id = input("请输入门诊号/住院号：")
    clinical_diagnosis = input("请输入临床诊断：")
    referring_doctor = input("请输入送检医生：")
    test_date = input("请输入检验日期：")
    report_date = input("请输入报告日期：")

    generate_report(name, gender, age, department, patient_id, clinical_diagnosis, referring_doctor, test_date, report_date)