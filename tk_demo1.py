import tkinter as tk
from tkinter import ttk
from docx import Document
import datetime
import os
import tempfile
import win32api
import win32print

def generate_report():
    doc = Document()
    doc.add_heading('检查报告', 0)
    # 头部信息
    doc.add_paragraph(f'姓名：{name_entry.get()}')
    doc.add_paragraph(f'性别：{gender_entry.get()}')
    doc.add_paragraph(f'年龄：{age_entry.get()}')
    doc.add_paragraph(f'标本类型：{specimen_type_entry.get()}')
    doc.add_paragraph(f'检查日期：{date_entry.get()}')
    doc.add_paragraph(f'科室：{department_entry.get()}')
    doc.add_paragraph(f'门诊号/住院号：{outpatient_entry.get()}')
    doc.add_paragraph(f'临床诊断：{diagnosis_entry.get()}')
    doc.add_paragraph(f'送检医生：{recommending_doctor_entry.get()}')
    # 检查结果
    doc.add_heading('检查结果', level=1)
    for row in range(8):
        doc.add_paragraph(f'序号：{row + 1}\n检查项目：\n结果：\n参考值：')
    # 底部信息
    doc.add_paragraph(f'检验时间：{inspection_time_entry.get()}')
    doc.add_paragraph(f'报告时间：{report_time_entry.get()}')
    doc.add_paragraph(f'检验者：{inspector_entry.get()}')
    doc.add_paragraph(f'审核者：{reviewer_entry.get()}')
    temp_file = tempfile.mktemp(suffix='.docx')
    doc.save(temp_file)
    return temp_file

def preview_report():
    temp_file = generate_report()
    os.startfile(temp_file)

def print_report():
    temp_file = generate_report()
    win32api.ShellExecute(0, "print", temp_file, None, ".", 0)

root = tk.Tk()
root.title('检查报告生成工具')

# 区域 1：头部信息区
header_frame = ttk.Frame(root)
header_frame.pack(pady=10)
name_label = ttk.Label(header_frame, text='姓名：')
name_label.grid(row=0, column=0, sticky='w')
name_entry = ttk.Entry(header_frame)
name_entry.grid(row=0, column=1)
gender_label = ttk.Label(header_frame, text='性别：')
gender_label.grid(row=0, column=2, sticky='w')
gender_entry = ttk.Entry(header_frame)
gender_entry.grid(row=0, column=3)
age_label = ttk.Label(header_frame, text='年龄：')
age_label.grid(row=0, column=4, sticky='w')
age_entry = ttk.Entry(header_frame)
age_entry.grid(row=0, column=5)
specimen_type_label = ttk.Label(header_frame, text='标本类型：')
specimen_type_label.grid(row=1, column=0, sticky='w')
specimen_type_entry = ttk.Entry(header_frame)
specimen_type_entry.grid(row=1, column=1)
date_label = ttk.Label(header_frame, text='检查日期：')
date_label.grid(row=1, column=2, sticky='w')
date_entry = ttk.Entry(header_frame)
date_entry.insert(0, datetime.date.today())
department_label = ttk.Label(header_frame, text='科室：')
department_label.grid(row=1, column=4, sticky='w')
department_entry = ttk.Entry(header_frame)
department_entry.grid(row=1, column=5)
outpatient_label = ttk.Label(header_frame, text='门诊号/住院号：')
outpatient_label.grid(row=2, column=0, sticky='w')
outpatient_entry = ttk.Entry(header_frame)
outpatient_entry.grid(row=2, column=1)
diagnosis_label = ttk.Label(header_frame, text='临床诊断：')
diagnosis_label.grid(row=2, column=2, sticky='w')
diagnosis_entry = ttk.Entry(header_frame)
diagnosis_entry.grid(row=2, column=3)
recommending_doctor_label = ttk.Label(header_frame, text='送检医生：')
recommending_doctor_label.grid(row=2, column=4, sticky='w')
recommending_doctor_entry = ttk.Entry(header_frame)
recommending_doctor_entry.grid(row=2, column=5)


# 区域 2：检查结果区
results_frame = ttk.Frame(root)
results_frame.pack(pady=10)
# style = ttk.Style()
# style.configure('Treeview', rowheight=25)
# tree = ttk.Treeview(results_frame, columns=('序号', '检查项目', '结果', '参考值'), show='headings')
# tree.heading('序号', text='序号')
# tree.heading('检查项目', text='检查项目')
# tree.heading('结果', text='结果')
# tree.heading('参考值', text='参考值')
# for _ in range(8):
#     tree.insert('', 'end', values=('', '', '', ''))
# tree.pack()
# tree.tag_configure('center', anchor='center')




# 区域 3：底部信息区
bottom_frame = ttk.Frame(root)
bottom_frame.pack(pady=10)
inspection_time_label = ttk.Label(bottom_frame, text='检验时间：')
inspection_time_label.grid(row=0, column=0, sticky='w')
inspection_time_entry = ttk.Entry(bottom_frame)
inspection_time_entry.grid(row=0, column=1)
report_time_label = ttk.Label(bottom_frame, text='报告时间：')
report_time_label.grid(row=0, column=2, sticky='w')
report_time_entry = ttk.Entry(bottom_frame)
report_time_entry.insert(0, datetime.date.today())
inspector_label = ttk.Label(bottom_frame, text='检验者：')
inspector_label.grid(row=1, column=0, sticky='w')
inspector_entry = ttk.Entry(bottom_frame)
inspector_entry.grid(row=1, column=1)
reviewer_label = ttk.Label(bottom_frame, text='审核者：')
reviewer_label.grid(row=1, column=2, sticky='w')
reviewer_entry = ttk.Entry(bottom_frame)
reviewer_entry.grid(row=1, column=3)

# 区域 4：按钮区
button_frame = ttk.Frame(root)
button_frame.pack(pady=10)
generate_button = ttk.Button(button_frame, text='生成报告', command=generate_report)
generate_button.grid(row=0, column=0)
preview_button = ttk.Button(button_frame, text='预览报告', command=preview_report)
preview_button.grid(row=0, column=1)
print_button = ttk.Button(button_frame, text='打印报告', command=print_report)
print_button.grid(row=0, column=2)

root.mainloop()