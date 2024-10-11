import tkinter as tk
from tkinter import ttk
from docx import Document
from datetime import datetime, date
import os
import time
import tempfile
from matplotlib.pyplot import grid
import win32api
import win32print

def generate_report():
    doc = Document()
    doc.add_heading('检查报告', 0)
    # 头部信息
    doc.add_paragraph(f'姓名：{name_entry.get()}')
    doc.add_paragraph(f'性别：{gender_entry.get()}')
    doc.add_paragraph(f'年龄：{age_entry.get()}')
    doc.add_paragraph(f'标本类型：{specimen_type_entry.get()}')
    doc.add_paragraph(f'检查日期：{date_entry.get()}')
    doc.add_paragraph(f'科室：{department_entry.get()}')
    doc.add_paragraph(f'门诊号/住院号：{outpatient_entry.get()}')
    doc.add_paragraph(f'临床诊断：{diagnosis_entry.get()}')
    doc.add_paragraph(f'送检医生：{recommending_doctor_entry.get()}')
    # 检查结果
    doc.add_heading('检查结果', level=1)
    for row in range(8):
        item = tree.item(tree.get_children()[row])
        values = item['values']
        doc.add_paragraph(f'序号：{values[0]}\n检查项目：{values[1]}\n结果：{values[2]}\n参考值：{values[3]}')
    # 底部信息
    doc.add_paragraph(f'检验时间：{inspection_time_entry.get()}')
    doc.add_paragraph(f'报告时间：{report_time_entry.get()}')
    doc.add_paragraph(f'检验者：{inspector_entry.get()}')
    doc.add_paragraph(f'审核者：{reviewer_entry.get()}')
    temp_file = tempfile.mktemp(suffix='.docx')
    doc.save(temp_file)
    return temp_file


def build_doc():
    # 基准文档
    doc = Document("HP1.docx")
    
    # 修改报告里的表格数据。
    today = date.today()
    date_now = today.strftime('%Y-%m-%d')
    
    tables = doc.tables
    
    # 处理第一行的姓名、性别、年龄、标本种类、送检日期
    table1 = tables[0]
    
    run = table1.cell(0, 0).paragraphs[0].add_run() # 姓名
    run.text = name_entry.get()
    
    run = table1.cell(0, 1).paragraphs[0].add_run() # 行别
    run.text = gender_entry.get()
    
    run = table1.cell(0, 2).paragraphs[0].add_run() # 年龄
    run.text = age_entry.get()
    
    run = table1.cell(0, 3).paragraphs[0].add_run() # 类型
    run.text = specimen_type_entry.get()
    
    run = table1.cell(0, 4).paragraphs[0].add_run() # 日期
    run.text = f"{date_now}"
    
    # 处理第二行的科室、门诊号/住院号、临床诊断、送检医生
    table2 = tables[1]
    
    run = table2.cell(0, 0).paragraphs[0].add_run() # 科室
    run.text = department_entry.get()
    
    run = table2.cell(0, 1).paragraphs[0].add_run() # 门诊号
    run.text = outpatient_entry.get()
    
    run = table2.cell(0, 2).paragraphs[0].add_run() # 诊断结果
    run.text = diagnosis_entry.get()
    
    run = table2.cell(0, 3).paragraphs[0].add_run() # 医生
    run.text = recommending_doctor_entry.get()
    
    
    # 处理检测结果序号、检查项目、结果、参考值
    table4 = tables[2]
    
    if check1_label.get() != "":
        run = table4.cell(1, 0).paragraphs[0].add_run() # 序号
        run.text = "1"
        run = table4.cell(1, 1).paragraphs[0].add_run() # 项目
        run.text = check1_label.get()
        run = table4.cell(1, 2).paragraphs[0].add_run() # 结果
        run.text = res1_label.get()
        run = table4.cell(1, 3).paragraphs[0].add_run() # 参考值
        run.text = "阴性（—）"
    
    if check2_label.get() != "":
        run = table4.cell(2, 0).paragraphs[0].add_run() # 序号
        run.text = "2"
        run = table4.cell(2, 1).paragraphs[0].add_run() # 项目
        run.text = check2_label.get()
        run = table4.cell(2, 2).paragraphs[0].add_run() # 结果
        run.text = res2_label.get()
        run = table4.cell(2, 3).paragraphs[0].add_run() # 参考值
        run.text = "阴性（—）"
    
    if check3_label.get() != "":
        run = table4.cell(3, 0).paragraphs[0].add_run() # 序号
        run.text = "3"
        run = table4.cell(3, 1).paragraphs[0].add_run() # 项目
        run.text = check3_label.get()
        run = table4.cell(3, 2).paragraphs[0].add_run() # 结果
        run.text = res3_label.get()
        run = table4.cell(3, 3).paragraphs[0].add_run() # 参考值
        run.text = "阴性（—）"
    
    if check4_label.get() != "":
        run = table4.cell(4, 0).paragraphs[0].add_run() # 序号
        run.text = "4"
        run = table4.cell(4, 1).paragraphs[0].add_run() # 项目
        run.text = check4_label.get()
        run = table4.cell(4, 2).paragraphs[0].add_run() # 结果
        run.text = res4_label.get()
        run = table4.cell(4, 3).paragraphs[0].add_run() # 参考值
        run.text = "阴性（—）"
    
    if check5_label.get() != "":
        run = table4.cell(5, 0).paragraphs[0].add_run() # 序号
        run.text = "5"
        run = table4.cell(5, 1).paragraphs[0].add_run() # 项目
        run.text = check5_label.get()
        run = table4.cell(5, 2).paragraphs[0].add_run() # 结果
        run.text = res5_label.get()
        run = table4.cell(5, 3).paragraphs[0].add_run() # 参考值
        run.text = "阴性（—）"
    
    if check6_label.get() != "":
        run = table4.cell(6, 0).paragraphs[0].add_run() # 序号
        run.text = "6"
        run = table4.cell(6, 1).paragraphs[0].add_run() # 项目
        run.text = check6_label.get()
        run = table4.cell(6, 2).paragraphs[0].add_run() # 结果
        run.text = res6_label.get()
        run = table4.cell(6, 3).paragraphs[0].add_run() # 参考值
        run.text = "阴性（—）"
    
    if check7_label.get() != "":
        run = table4.cell(7, 0).paragraphs[0].add_run() # 序号
        run.text = "7"
        run = table4.cell(7, 1).paragraphs[0].add_run() # 项目
        run.text = check7_label.get()
        run = table4.cell(7, 2).paragraphs[0].add_run() # 结果
        run.text = res7_label.get()
        run = table4.cell(7, 3).paragraphs[0].add_run() # 参考值
        run.text = "阴性（—）"
    
    if check8_label.get() != "":
        run = table4.cell(8, 0).paragraphs[0].add_run() # 序号
        run.text = "8"
        run = table4.cell(8, 1).paragraphs[0].add_run() # 项目
        run.text = check8_label.get()
        run = table4.cell(8, 2).paragraphs[0].add_run() # 结果
        run.text = res8_label.get()
        run = table4.cell(8, 3).paragraphs[0].add_run() # 参考值
        run.text = "阴性（—）"
    
    
    # 处理底部的检验时间、报告时间、检验者、审核者
    # table_list[1].rows[0].cells[0].text   '检验时间：                     报告时间：            检验者：       审核者：                 '
    table_l = tables[-1]
    
    run = table_l.cell(0, 0).paragraphs[0].add_run()
    run.text = f"{date_now}"
    run = table_l.cell(0, 1).paragraphs[0].add_run()
    run.text = f"{date_now}"
    
    run = table_l.cell(0, 2).paragraphs[0].add_run()
    run.text = inspector_entry.get()    # 检验者
    run = table_l.cell(0, 3).paragraphs[0].add_run()
    run.text = reviewer_entry.get()  # 审核者
    
    # 保存word文档
    doc.save(f"test_{datetime.today().strftime('%Y%m%d-%H%M%S')}.docx")


def preview_report():
    temp_file = generate_report()
    os.startfile(temp_file)

def print_report():
    temp_file = generate_report()
    win32api.ShellExecute(0, "print", temp_file, None, ".", 0)

root = tk.Tk()
root.title('检查报告生成工具')

# 定义性别变量
gender_choice = tk.StringVar()
docker_var = tk.StringVar()

# 区域 1：头部信息区
header_frame = ttk.Frame(root)
header_frame.pack(pady=10)

name_label = ttk.Label(header_frame, text='姓名：')
name_label.grid(row=0, column=0, sticky='w')
name_entry = ttk.Entry(header_frame)
name_entry.grid(row=0, column=1)

gender_label = ttk.Label(header_frame, text='性别：')
gender_label.grid(row=0, column=2, sticky='w')
# gender_entry = ttk.Entry(header_frame)
gender_entry = ttk.Combobox(header_frame, textvariable=gender_choice, values=['男', '女'])
gender_entry.grid(row=0, column=3)

age_label = ttk.Label(header_frame, text='年龄：')
age_label.grid(row=0, column=4, sticky='w')
age_entry = ttk.Entry(header_frame)
age_entry.grid(row=0, column=5)

specimen_type_label = ttk.Label(header_frame, text='标本类型：')
specimen_type_label.grid(row=1, column=0, sticky='w')
specimen_type_entry = ttk.Entry(header_frame)
specimen_type_entry.grid(row=1, column=1)
specimen_type_entry.insert(0, '末梢血')

date_label = ttk.Label(header_frame, text='检查日期：')
date_label.grid(row=1, column=2, sticky='w')
date_entry = ttk.Entry(header_frame)
date_entry.grid(row=1, column=3, sticky='w')
date_entry.insert(0, date.today())

department_label = ttk.Label(header_frame, text='科室：')
department_label.grid(row=1, column=4, sticky='w')
department_entry = ttk.Entry(header_frame)
department_entry.grid(row=1, column=5)
department_entry.insert(0, '消化内科')

outpatient_label = ttk.Label(header_frame, text='门诊号/住院号：')
outpatient_label.grid(row=2, column=0, sticky='w')
outpatient_entry = ttk.Entry(header_frame)
outpatient_entry.grid(row=2, column=1)

diagnosis_label = ttk.Label(header_frame, text='临床诊断：')
diagnosis_label.grid(row=2, column=2, sticky='w')
diagnosis_entry = ttk.Entry(header_frame)
diagnosis_entry.grid(row=2, column=3)

recommending_doctor_label = ttk.Label(header_frame, text='送检医生：')
recommending_doctor_label.grid(row=2, column=4, sticky='w')
# recommending_doctor_entry = ttk.Entry(header_frame)
recommending_doctor_entry = ttk.Combobox(header_frame, textvariable=docker_var, values=['陈光耀', '陈颖', '李峰', '祝子华', '万红宇', '刘海玲', '方青青', '钱孝先', '沈丹杰', '陈炜', '张君佩', '赵娟', '吴冰', '龚鹏', '姚灿', '李煜', '徐鑫鑫', '李晓娟'])
recommending_doctor_entry.grid(row=2, column=5)

# 区域 2：检查结果区
results_frame = ttk.Frame(root)
results_frame.pack(pady=10)
# style = ttk.Style()

# 定义 8 个结果变量
resutls1_yin_yang = tk.StringVar()
resutls2_yin_yang = tk.StringVar()
resutls3_yin_yang = tk.StringVar()
resutls4_yin_yang = tk.StringVar()
resutls5_yin_yang = tk.StringVar()
resutls6_yin_yang = tk.StringVar()
resutls7_yin_yang = tk.StringVar()
resutls8_yin_yang = tk.StringVar()


num_label = ttk.Label(results_frame, text='序号').grid(row=4, column=1, columnspan=2, sticky="w", padx=10, pady=5)
check_label = ttk.Label(results_frame, text='检查项目').grid(row=4, column=5, columnspan=10, sticky="w", padx=10, pady=5)
res_label = ttk.Label(results_frame, text='结果').grid(row=4, column=17, columnspan=2, sticky="w", padx=10, pady=5)
reference_label = ttk.Label(results_frame, text='参考值').grid(row=4, column=21, columnspan=2, sticky="w", padx=10, pady=5)

num1_label = ttk.Label(results_frame, text='1').grid(row=6, column=1, columnspan=2, sticky="w", padx=10, pady=5)
check1_label = ttk.Entry(results_frame, width=40).grid(row=6, column=5, columnspan=10, sticky="w", padx=10, pady=5)
# res1_label = ttk.Entry(results_frame).grid(row=6, column=17, columnspan=2, sticky="w", padx=10, pady=5)
res1_label = ttk.Combobox(results_frame, width=10, textvariable=resutls1_yin_yang, values=['阴性（—）', '阳性（—）']).grid(row=6, column=17, columnspan=2, sticky="w", padx=10, pady=5)
reference1_label = ttk.Label(results_frame, text='阴性（—）').grid(row=6, column=21, columnspan=2, sticky="w", padx=10, pady=5)

num2_label = ttk.Label(results_frame, text='2').grid(row=8, column=1, columnspan=2, sticky="w", padx=10, pady=5)
check2_label = ttk.Entry(results_frame, width=40).grid(row=8, column=5, columnspan=10, sticky="w", padx=10, pady=5)
# res2_label = ttk.Entry(results_frame).grid(row=8, column=17, columnspan=2, sticky="w", padx=10, pady=5)
res2_label = ttk.Combobox(results_frame, width=10, textvariable=resutls2_yin_yang, values=['阴性（—）', '阳性（—）']).grid(row=8, column=17, columnspan=2, sticky="w", padx=10, pady=5)
reference2_label = ttk.Label(results_frame, text='阴性（—）').grid(row=8, column=21, columnspan=2, sticky="w", padx=10, pady=5)

num3_label = ttk.Label(results_frame, text='3').grid(row=10, column=1, columnspan=2, sticky="w", padx=10, pady=5)
check3_label = ttk.Entry(results_frame, width=40).grid(row=10, column=5, columnspan=10, sticky="w", padx=10, pady=5)
# res3_label = ttk.Entry(results_frame).grid(row=10, column=17, columnspan=2, sticky="w", padx=10, pady=5)
res3_label = ttk.Combobox(results_frame, width=10, textvariable=resutls3_yin_yang, values=['阴性（—）', '阳性（—）']).grid(row=10, column=17, columnspan=2, sticky="w", padx=10, pady=5)
reference3_label = ttk.Label(results_frame, text='阴性（—）').grid(row=10, column=21, columnspan=2, sticky="w", padx=10, pady=5)

num4_label = ttk.Label(results_frame, text='4').grid(row=12, column=1, columnspan=2, sticky="w", padx=10, pady=5)
check4_label = ttk.Entry(results_frame, width=40).grid(row=12, column=5, columnspan=10, sticky="w", padx=10, pady=5)
# res4_label = ttk.Entry(results_frame).grid(row=12, column=17, columnspan=2, sticky="w", padx=10, pady=5)
res4_label = ttk.Combobox(results_frame, width=10, textvariable=resutls4_yin_yang, values=['阴性（—）', '阳性（—）']).grid(row=12, column=17, columnspan=2, sticky="w", padx=10, pady=5)
reference4_label = ttk.Label(results_frame, text='阴性（—）').grid(row=12, column=21, columnspan=2, sticky="w", padx=10, pady=5)

num5_label = ttk.Label(results_frame, text='5').grid(row=14, column=1, columnspan=2, sticky="w", padx=10, pady=5)
check5_label = ttk.Entry(results_frame, width=40).grid(row=14, column=5, columnspan=10, sticky="w", padx=10, pady=5)
# res5_label = ttk.Entry(results_frame).grid(row=14, column=17, columnspan=2, sticky="w", padx=10, pady=5)
res5_label = ttk.Combobox(results_frame, width=10, textvariable=resutls5_yin_yang, values=['阴性（—）', '阳性（—）']).grid(row=14, column=17, columnspan=2, sticky="w", padx=10, pady=5)
reference5_label = ttk.Label(results_frame, text='阴性（—）').grid(row=14, column=21, columnspan=2, sticky="w", padx=10, pady=5)

num6_label = ttk.Label(results_frame, text='6').grid(row=16, column=1, columnspan=2, sticky="w", padx=10, pady=5)
check6_label = ttk.Entry(results_frame, width=40).grid(row=16, column=5, columnspan=10, sticky="w", padx=10, pady=5)
# res6_label = ttk.Entry(results_frame).grid(row=16, column=17, columnspan=2, sticky="w", padx=10, pady=5)
res6_label = ttk.Combobox(results_frame, width=10, textvariable=resutls6_yin_yang, values=['阴性（—）', '阳性（—）']).grid(row=16, column=17, columnspan=2, sticky="w", padx=10, pady=5)
reference6_label = ttk.Label(results_frame, text='阴性（—）').grid(row=16, column=21, columnspan=2, sticky="w", padx=10, pady=5)

num7_label = ttk.Label(results_frame, text='7').grid(row=18, column=1, columnspan=2, sticky="w", padx=10, pady=5)
check7_label = ttk.Entry(results_frame, width=40).grid(row=18, column=5, columnspan=10, sticky="w", padx=10, pady=5)
# res7_label = ttk.Entry(results_frame).grid(row=18, column=17, columnspan=2, sticky="w", padx=10, pady=5)
res7_label = ttk.Combobox(results_frame, width=10, textvariable=resutls7_yin_yang, values=['阴性（—）', '阳性（—）']).grid(row=18, column=17, columnspan=2, sticky="w", padx=10, pady=5)
reference7_label = ttk.Label(results_frame, text='阴性（—）').grid(row=18, column=21, columnspan=2, sticky="w", padx=10, pady=5)

num8_label = ttk.Label(results_frame, text='8').grid(row=20, column=1, columnspan=2, sticky="w", padx=10, pady=5)
check8_label = ttk.Entry(results_frame, width=40).grid(row=20, column=5, columnspan=10, sticky="w", padx=10, pady=5)
# res8_label = ttk.Entry(results_frame).grid(row=20, column=17, columnspan=2, sticky="w", padx=10, pady=5)
res8_label = ttk.Combobox(results_frame, width=10, textvariable=resutls8_yin_yang, values=['阴性（—）', '阳性（—）']).grid(row=20, column=17, columnspan=2, sticky="w", padx=10, pady=5)
reference8_label = ttk.Label(results_frame, text='阴性（—）').grid(row=20, column=21, columnspan=2, sticky="w", padx=10, pady=5)



# 区域 3：底部信息区
bottom_frame = ttk.Frame(root)
bottom_frame.pack(pady=10)

inspection_time_label = ttk.Label(bottom_frame, text='检验时间：')
inspection_time_label.grid(row=0, column=0, sticky='w')
inspection_time_entry = ttk.Entry(bottom_frame)
inspection_time_entry.grid(row=0, column=1)
inspection_time_entry.insert(0, date.today())

report_time_label = ttk.Label(bottom_frame, text='报告时间：')
report_time_label.grid(row=0, column=2, sticky='w')
report_time_entry = ttk.Entry(bottom_frame)
report_time_entry.grid(row=0, column=3)
report_time_entry.insert(0, date.today())

inspector_label = ttk.Label(bottom_frame, text='检验者：')
inspector_label.grid(row=1, column=0, sticky='w')
inspector_entry = ttk.Entry(bottom_frame)
inspector_entry.grid(row=1, column=1)

reviewer_label = ttk.Label(bottom_frame, text='审核者：')
reviewer_label.grid(row=1, column=2, sticky='w')
reviewer_entry = ttk.Entry(bottom_frame)
reviewer_entry.grid(row=1, column=3)



# 区域 4：按钮区
button_frame = ttk.Frame(root)
button_frame.pack(pady=10)
generate_button = ttk.Button(button_frame, text='生成报告', command=build_doc)
generate_button.grid(row=0, column=0)
preview_button = ttk.Button(button_frame, text='预览报告', command=preview_report)
preview_button.grid(row=0, column=1)
print_button = ttk.Button(button_frame, text='打印报告', command=print_report)
print_button.grid(row=0, column=2)

root.mainloop()